import sys
import uuid
import yaml

import docx
import qrcode

d = docx.Document('/vagrant/obtig-user-instructions.docx')
msl = d.tables[1].rows[0].cells[1].paragraphs[0].style  # monospaced left
msc = d.tables[-1].rows[-1].cells[-1].paragraphs[0].style  # monospaced center

cluster_number = int(sys.argv[1])

with open(f"/vagrant/ansible/host_vars/hpc{cluster_number}-sms", "r") as f:
    y = yaml.load(f, yaml.CLoader)
with open("/vagrant/ansible/user-passwords.txt", "r") as f:
    p = f.readlines()

password = p[cluster_number].strip()
sms_public_ip = y['sms_ipv4']
sms_external_ip = f"10.38.51.{cluster_number}"
login_public_ip = y['login_ipv4']
login_external_ip = f"10.38.51.{cluster_number+128}"
login_mac = y['login_mac']
c1_mac = y['compute_nodes'][0]['mac']
c2_mac = y['compute_nodes'][1]['mac']
#g1_mac = y['gpu_nodes'][0]['mac']
#g2_mac = y['gpu_nodes'][1]['mac']

# Title of second page
d.paragraphs[8].text = d.paragraphs[8].text.replace('N', str(cluster_number))

# QR code
filename = f"student-{cluster_number}-{uuid.uuid4()}"
qr = qrcode.QRCode()
qr.add_data(f'https://www.hpc.tntech.edu/obtig/{filename}.pdf')
qr.make(fit=True)
img = qr.make_image()
img.save("qr.png")
qr_run = d.tables[0].rows[1].cells[1].paragraphs[0].add_run()
qr_run.add_picture('qr.png', width=docx.shared.Inches(1.25))

# Credentials table
d.tables[1].rows[1].cells[1].text = password
d.tables[1].rows[1].cells[1].paragraphs[0].style = msl

# SMS table
d.tables[3].rows[1].cells[1].text = sms_public_ip
d.tables[3].rows[1].cells[1].paragraphs[0].style = msl
d.tables[3].rows[2].cells[1].text = sms_external_ip
d.tables[3].rows[2].cells[1].paragraphs[0].style = msl

# Login table
d.tables[4].rows[1].cells[1].text = login_public_ip
d.tables[4].rows[1].cells[1].paragraphs[0].style = msl
d.tables[4].rows[2].cells[1].text = login_external_ip
d.tables[4].rows[2].cells[1].paragraphs[0].style = msl
d.tables[4].rows[4].cells[1].text = login_mac
d.tables[4].rows[4].cells[1].paragraphs[0].style = msl

# Compute nodes table
d.tables[5].rows[1].cells[1].text = c1_mac
d.tables[5].rows[1].cells[1].paragraphs[0].style = msc
d.tables[5].rows[2].cells[1].text = c2_mac
d.tables[5].rows[2].cells[1].paragraphs[0].style = msc
#d.tables[5].rows[3].cells[1].text = g1_mac
#d.tables[5].rows[3].cells[1].paragraphs[0].style = msc
#d.tables[5].rows[4].cells[1].text = g2_mac
#d.tables[5].rows[4].cells[1].paragraphs[0].style = msc

d.save(f"/vagrant/{filename}.docx")
